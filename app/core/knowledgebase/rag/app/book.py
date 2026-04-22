"""
书籍文档切片方法
支持PDF、DOCX等格式的书籍结构化解析和章节切片
"""

import copy
import re
import logging
from io import BytesIO

logger = logging.getLogger(__name__)


def chunk(filename, binary=None, from_page=0, to_page=100000, lang="Chinese", callback=None, **kwargs):
    """
    书籍切片函数
    
    支持PDF、DOCX格式，按章节结构进行切片
    
    Args:
        filename: 文件名或文件路径
        binary: 文件二进制数据（可选）
        from_page: 起始页码
        to_page: 结束页码
        lang: 语言 ("Chinese" 或 "English")
        callback: 进度回调函数 callback(progress, message)
        **kwargs: 其他参数
        
    Returns:
        list: 切片后的文档列表
    """
    from ..nlp import (
        naive_merge,
        tokenize_chunks,
        rag_tokenizer,
        num_tokens_from_string,
        is_english,
        tokenize_doc,
    )
    
    parser_config = kwargs.get("parser_config", {
        "chunk_token_num": 512,
        "delimiter": "\n。；！？",
        "layout_recognize": "DeepDOC",
    })
    
    eng = lang.lower() == "english"
    doc = {"docnm_kwd": filename}
    doc["title_tks"] = rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", filename))
    doc["title_sm_tks"] = rag_tokenizer.fine_grained_tokenize(doc["title_tks"])
    res = []
    
    # PDF文件处理
    if re.search(r"\.pdf$", filename, re.IGNORECASE):
        callback(0.1, "开始解析PDF书籍")
        
        sections, tables, pdf_parser = _parse_pdf(
            filename=filename,
            binary=binary,
            from_page=from_page,
            to_page=to_page,
            lang=lang,
            callback=callback,
            **kwargs
        )
        
        if not sections and not tables:
            return []
            
        res.extend(_process_book_sections(sections, doc, eng, parser_config))
        callback(0.8, "PDF书籍解析完成")
        
    # DOCX文件处理
    elif re.search(r"\.docx$", filename, re.IGNORECASE):
        callback(0.1, "开始解析DOCX书籍")
        
        sections = _parse_docx(filename, binary)
        
        # 按章节标题分割
        chapter_sections = _split_by_chapters(sections)
        res.extend(_process_book_sections(chapter_sections, doc, eng, parser_config))
        callback(0.8, "DOCX书籍解析完成")
        
    else:
        raise NotImplementedError(f"不支持的文件类型: {filename} (支持: pdf, docx)")
    
    return res


def _parse_pdf(filename, binary=None, from_page=0, to_page=100000, callback=None, **kwargs):
    """解析PDF文件"""
    try:
        from app.core.knowledgebase.deepdoc.parser import PdfParser
        
        pdf_parser = PdfParser()
        sections, tables = pdf_parser(
            filename if not binary else binary, 
            from_page=from_page, 
            to_page=to_page, 
            callback=callback
        )
        return sections, tables, pdf_parser
    except Exception as e:
        logger.error(f"Failed to parse PDF book: {e}")
        return [], [], None


def _parse_docx(filename, binary=None):
    """解析DOCX文件"""
    try:
        from docx import Document
        
        doc = Document(filename) if not binary else Document(BytesIO(binary))
        sections = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                style_name = para.style.name if para.style else ""
                is_heading = 'Heading' in style_name or 'heading' in style_name.lower() or para.style.name.startswith('Heading')
                sections.append((text, is_heading))
                
        return sections
    except Exception as e:
        logger.error(f"Failed to parse DOCX book: {e}")
        return []


def _split_by_chapters(sections):
    """按章节标题分割内容"""
    chapters = []
    current_chapter = []
    
    for text, is_heading in sections:
        if is_heading and current_chapter:
            chapters.append(current_chapter)
            current_chapter = [text]
        else:
            current_chapter.append(text)
            
    if current_chapter:
        chapters.append(current_chapter)
        
    return chapters


def _process_book_sections(sections, doc, eng, parser_config):
    """处理书籍章节并生成chunks"""
    from ..nlp import tokenize_chunks
    
    chunks = []
    chunk_token_num = int(parser_config.get("chunk_token_num", 512))
    delimiter = parser_config.get("delimiter", "\n")
    
    for section in sections:
        if isinstance(section, list):
            content = "\n".join(str(s) for s in section if s)
        elif isinstance(section, tuple):
            content = section[0]
        else:
            content = str(section)
            
        if not content.strip():
            continue
            
        # 使用naive merge进行分块
        merged_chunks = _book_naive_merge(content, chunk_token_num, delimiter)
        chunks.extend(merged_chunks)
    
    # 对所有chunks进行tokenize处理
    res = []
    for ck in chunks:
        d = copy.deepcopy(doc)
        tokenize_doc(d, ck, eng)
        res.append(d)
        
    return res


def _book_naive_merge(text, chunk_token_num=512, delimiter="\n"):
    """
    书籍专用文本合并算法
    
    优先在章节边界处切分，保持章节完整性
    """
    if not text or not text.strip():
        return []
    
    # 按分隔符分割
    custom_delimiters = [m.group(1) for m in re.finditer(r"`([^`]+)`", delimiter)]
    has_custom = bool(custom_delimiters)
    
    if has_custom:
        custom_pattern = "|".join(re.escape(t) for t in sorted(set(custom_delimiters), key=len, reverse=True))
        parts = re.split(r"(%s)" % custom_pattern, text, flags=re.DOTALL)
        parts = [p for p in parts if p and not re.fullmatch(custom_pattern, p)]
    else:
        parts = re.split(r"[%s]" % re.escape(delimiter), text)
    
    cks = []
    current_chunk = ""
    current_tokens = 0
    
    for part in parts:
        part_tokens = num_tokens_from_string(part)
        
        if current_tokens + part_tokens > chunk_token_num and current_chunk:
            cks.append(current_chunk.strip())
            overlap_len = max(0, len(current_chunk) // 4)
            current_chunk = current_chunk[-overlap_len:] if overlap_len > 0 else ""
            current_tokens = num_tokens_from_string(current_chunk)
        
        current_chunk += (("\n" + part) if current_chunk else part)
        current_tokens += part_tokens
    
    if current_chunk.strip():
        cks.append(current_chunk.strip())
    
    return cks


if __name__ == "__main__":
    import sys
    
    def dummy_callback(prog=None, msg=""):
        print(f"[{prog}] {msg}" if prog else msg)
    
    if len(sys.argv) > 1:
        result = chunk(sys.argv[1], callback=dummy_callback)
        print(f"Total chunks: {len(result)}")
