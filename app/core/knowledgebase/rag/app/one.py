"""
单一文档切片方法
将整个文档作为一个chunk，保持原始文本顺序
"""

import copy
import re
import logging
from io import BytesIO

logger = logging.getLogger(__name__)


class Pdf:
    """PDF解析器（用于one模式）"""
    
    def __call__(self, filename, binary=None, from_page=0, to_page=100000, zoomin=3, callback=None):
        """解析PDF并保持原始文本顺序"""
        try:
            from app.core.knowledgebase.deepdoc.parser import PdfParser
            
            pdf_parser = PdfParser()
            
            # 1. OCR
            callback(msg="开始OCR识别")
            pdf_parser.__images__(filename if not binary else binary, zoomin, from_page, to_page, callback)
            
            # 2. 布局分析
            callback(msg="布局分析")
            pdf_parser._layouts_rec(zoomin, drop=False)
            
            # 3. 表格分析
            callback(msg="表格分析")
            pdf_parser._table_transformer_job(zoomin)
            
            # 4. 文本合并
            pdf_parser._text_merge()
            
            # 5. 提取表格
            tbls = pdf_parser._extract_table_figure(True, zoomin, True, True)
            
            # 6. 向下合并
            pdf_parser._concat_downward()
            
            # 7. 提取所有文本框（保持顺序）
            sections = [
                (b["text"], pdf_parser.get_position(b, zoomin)) 
                for i, b in enumerate(pdf_parser.boxes)
            ]
            
            # 按位置排序
            sections.sort(key=lambda x: (x[-1][0][0], x[-1][0][3], x[-1][0][1]))
            
            return [(txt, "") for txt, _ in sections], tbls
            
        except Exception as e:
            logger.error(f"One模式PDF解析失败: {e}")
            return [], []


def chunk(filename, binary=None, from_page=0, to_page=100000, lang="Chinese", callback=None, **kwargs):
    """
    单一文档切片函数
    
    将整个文档作为一个chunk，保持原始文本顺序
    
    支持的文件格式: DOCX, PDF, Excel, TXT, Markdown, HTML, DOC
    
    Args:
        filename: 文件名或文件路径
        binary: 文件二进制数据（可选）
        from_page: 起始页码
        to_page: 结束页码
        lang: 语言 ("Chinese" 或 "English")
        callback: 进度回调函数 callback(progress, message)
        **kwargs: 其他参数
        
    Returns:
        list: 包含单个文档元素的列表
    """
    from ..nlp import rag_tokenizer, tokenize_doc
    
    parser_config = kwargs.get("parser_config", {
        "chunk_token_num": 512,
        "delimiter": "\n!?。；！？",
        "layout_recognize": "DeepDOC",
    })
    
    eng = lang.lower() == "english"
    sections = []
    
    # DOCX文件处理
    if re.search(r"\.docx$", filename, re.IGNORECASE):
        callback(0.1, "开始解析DOCX")
        sections = _parse_docx(filename, binary)
        callback(0.8, "DOCX解析完成")
        
    # PDF文件处理
    elif re.search(r"\.pdf$", filename, re.IGNORECASE):
        callback(0.1, "开始解析PDF")
        
        layout_recognizer = parser_config.get("layout_recognize", "DeepDOC")
        if isinstance(layout_recognizer, bool):
            layout_recognizer = "DeepDOC" if layout_recognizer else "Plain Text"
        
        name = layout_recognizer.strip().lower()
        
        pdf_parser_cls = Pdf
        sections, tbls, pdf_parser_obj = _parse_pdf(
            filename=filename,
            binary=binary,
            from_page=from_page,
            to_page=to_page,
            lang=lang,
            callback=callback,
            pdf_parser_cls=pdf_parser_cls,
            layout_recognizer=layout_recognizer,
            **kwargs
        )
        
        if not sections and not tbls:
            return []
        
        # 将表格内容也加入sections
        for (img, rows), poss in tbls:
            if rows:
                if isinstance(rows, str):
                    content = rows
                elif isinstance(rows, list):
                    content = rows[0] if rows else ""
                else:
                    content = str(rows)
                
                if poss:
                    adjusted_pos = [(p[0] + 1 - from_page, p[1], p[2], p[3], p[4]) for p in poss]
                    sections.append((content, adjusted_pos))
                else:
                    sections.append((content, ""))
        
        # 过滤空内容
        sections = [s for s, _ in sections if s]
        callback(0.8, "PDF解析完成")
        
    # Excel文件处理
    elif re.search(r"\.xlsx?$", filename, re.IGNORECASE):
        callback(0.1, "开始解析Excel")
        excel_sections = _parse_excel(binary)
        sections = excel_sections
        callback(0.8, "Excel解析完成")
        
    # Markdown/TXT文件处理
    elif re.search(r"\.(txt|md|markdown|mdx)$", filename, re.IGNORECASE):
        callback(0.1, "开始解析文本文件")
        text = _read_text_file(filename, binary)
        sections = text.split("\n")
        sections = [s for s in sections if s]
        callback(0.8, "文本文件解析完成")
        
    # HTML文件处理
    elif re.search(r"\.(htm|html)$", filename, re.IGNORECASE):
        callback(0.1, "开始解析HTML")
        html_sections = _parse_html(filename, binary)
        sections = [s for s in html_sections if s]
        callback(0.8, "HTML解析完成")
        
    # DOC文件处理（旧版Word）
    elif re.search(r"\.doc$", filename, re.IGNORECASE):
        callback(0.1, "开始解析DOC")
        doc_sections = _parse_doc_old(binary)
        sections = doc_sections
        callback(0.8, "DOC解析完成")
        
    else:
        raise NotImplementedError(f"不支持的文件类型: {filename} (支持: doc, docx, pdf, txt, md, html, xlsx)")
    
    # 构建单一文档
    doc = {
        "docnm_kwd": filename,
        "title_tks": rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", filename)),
    }
    doc["title_sm_tks"] = rag_tokenizer.fine_grained_tokenize(doc["title_tks"])
    
    # 合并所有内容为一个chunk
    full_text = "\n".join(str(s) for s in sections if s)
    tokenize_doc(doc, full_text, eng)
    
    return [doc]


def _parse_docx(filename, binary=None):
    """解析DOCX文件"""
    try:
        from docx import Document
        
        doc = Document(filename) if not binary else Document(BytesIO(binary))
        sections = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                sections.append(text)
                
        # 也提取表格
        for table in doc.tables:
            table_html = "<table>"
            for row in table.rows:
                table_html += "<tr>"
                for cell in row.cells:
                    table_html += f"<td>{cell.text}</td>"
                table_html += "</tr>"
            table_html += "</table>"
            sections.append(table_html)
        
        return sections
    except Exception as e:
        logger.error(f"DOCX解析失败: {e}")
        return []


def _parse_pdf(filename, binary=None, from_page=0, to_page=100000, lang="Chinese", 
               callback=None, pdf_parser_cls=None, layout_recognizer="DeepDOC", **kwargs):
    """解析PDF文件"""
    try:
        if pdf_parser_cls:
            pdf_parser = pdf_parser_cls()
            result = pdf_parser(
                filename=filename,
                binary=binary,
                from_page=from_page,
                to_page=to_page,
                callback=callback
            )
            if isinstance(result, tuple) and len(result) >= 2:
                sections, tables = result[0], result[1]
                return sections, tables, pdf_parser
        else:
            from app.core.knowledgebase.deepdoc.parser import PdfParser
            pdf_parser = PdfParser()
            sections, tables = pdf_parser(
                filename if not binary else binary, 
                from_page=from_page, 
                to_page=to_page, 
                callback=callback
            )
            return sections, tables, pdf_parser
    except Exception as e:
        logger.error(f"PDF解析失败: {e}")
        return [], [], None


def _parse_excel(binary=None):
    """解析Excel文件"""
    try:
        from openpyxl import load_workbook
        from io import BytesIO
    except ImportError:
        return []
    
    if binary:
        wb = load_workbook(BytesIO(binary), data_only=True)
    else:
        return []
    
    sections = []
    
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        for row in ws.iter_rows(values_only=True):
            row_text = "\t".join(str(cell) if cell is not None else "" for cell in row)
            if row_text.strip():
                sections.append(row_text)
    
    return sections


def _read_text_file(filename, binary=None):
    """读取文本文件"""
    from ..nlp import find_codec
    
    if binary:
        encoding = find_codec(binary)
        return binary.decode(encoding, errors="ignore")
    else:
        with open(filename, 'rb') as f:
            blob = f.read()
        encoding = find_codec(blob)
        return blob.decode(encoding, errors="ignore")


def _parse_html(filename, binary=None):
    """解析HTML文件"""
    try:
        from app.core.knowledgebase.deepdoc.parser import HtmlParser
        html_parser = HtmlParser()
        return html_parser(filename if not binary else binary, chunk_token_num=512)
    except Exception as e:
        logger.error(f"HTML解析失败: {e}")
        return []


def _parse_doc_old(binary=None):
    """解析旧版DOC文件"""
    try:
        from tika import parser as tika_parser
    except ImportError:
        logger.warning("tika不可用，无法解析.doc文件")
        return []
    
    if binary:
        data = BytesIO(binary)
    else:
        return []
    
    doc_parsed = tika_parser.from_buffer(data)
    if doc_parsed.get("content") is not None:
        return doc_parsed["content"].split("\n")
    return []


if __name__ == "__main__":
    import sys
    
    def dummy_callback(prog=None, msg=""):
        print(f"[{prog}] {msg}" if prog else msg)
    
    if len(sys.argv) > 1:
        result = chunk(sys.argv[1], callback=dummy_callback)
        print(f"Total chunks: {len(result)}")
