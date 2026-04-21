"""
论文文档切片方法
支持学术论文的结构化解析和多列布局处理
"""

import copy
import re
import logging
from collections import defaultdict
from io import BytesIO

try:
    from PIL import Image
except ImportError:
    Image = None

logger = logging.getLogger(__name__)


def chunk(filename, binary=None, from_page=0, to_page=100000, lang="Chinese", callback=None, **kwargs):
    """
    论文切片函数
    
    支持PDF、DOCX格式的论文解析，识别论文结构（标题、摘要、正文、参考文献等）
    
    Args:
        filename: 文件名或文件路径
        binary: 文件二进制数据（可选）
        from_page: 起始页码
        to_page: 结束页码
        lang: 语言 ("Chinese" 或 "English")
        callback: 进度回调函数 callback(progress, message)
        **kwargs: 其他参数
        
    Returns:
        list: 切片后的文档列表
    """
    from ..nlp import (
        naive_merge,
        tokenize_chunks,
        rag_tokenizer,
        tokenize,
        is_english,
    )
    
    parser_config = kwargs.get("parser_config", {
        "chunk_token_num": 1024,
        "delimiter": "\n。；！？",
        "layout_recognize": "DeepDOC",
    })
    
    eng = lang.lower() == "english"
    doc = {"docnm_kwd": filename}
    doc["title_tks"] = rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", filename))
    doc["title_sm_tks"] = rag_tokenizer.fine_grained_tokenize(doc["title_tks"])
    res = []
    
    # PDF文件处理
    if re.search(r"\.pdf$", filename, re.IGNORECASE):
        callback(0.1, "开始解析PDF论文")
        
        sections, tables, pdf_parser = _parse_paper_pdf(
            filename=filename,
            binary=binary,
            from_page=from_page,
            to_page=to_page,
            lang=lang,
            callback=callback,
            **kwargs
        )
        
        if not sections and not tables:
            return []
        
        # 识别论文结构
        structured_sections = _identify_paper_structure(sections)
        
        # 按结构化部分生成chunks
        for section_type, section_content in structured_sections:
            d = copy.deepcopy(doc)
            d["paper_section_kwd"] = section_type
            
            if isinstance(section_content, list):
                content = "\n".join(section_content)
            else:
                content = str(section_content)
                
            tokenize(d, content, eng)
            res.append(d)
        
        callback(0.8, "PDF论文解析完成")
        
    # DOCX文件处理
    elif re.search(r"\.docx$", filename, re.IGNORECASE):
        callback(0.1, "开始解析DOCX论文")
        
        sections = _parse_docx(filename, binary)
        structured_sections = _identify_paper_structure_from_docx(sections)
        
        for section_type, section_content in structured_sections:
            d = copy.deepcopy(doc)
            d["paper_section_kwd"] = section_type
            
            if isinstance(section_content, list):
                content = "\n".join(section_content)
            else:
                content = str(section_content)
                
            tokenize(d, content, eng)
            res.append(d)
        
        callback(0.8, "DOCX论文解析完成")
        
    else:
        raise NotImplementedError(f"不支持的文件类型: {filename} (支持: pdf, docx)")
    
    return res


def _parse_paper_pdf(filename, binary=None, from_page=0, to_page=100000, callback=None, **kwargs):
    """解析PDF论文"""
    try:
        from app.core.knowledgebase.deepdoc.parser import PdfParser
        
        pdf_parser = PdfParser()
        sections, tables = pdf_parser(
            filename if not binary else binary, 
            from_page=from_page, 
            to_page=to_page, 
            callback=callback
        )
        return sections, tables, pdf_parser
    except Exception as e:
        logger.error(f"Failed to parse PDF paper: {e}")
        return [], [], None


def _parse_docx(filename, binary=None):
    """解析DOCX论文"""
    try:
        from docx import Document
        
        doc = Document(filename) if not binary else Document(BytesIO(binary))
        sections = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                style_name = para.style.name if para.style else ""
                sections.append({
                    "text": text,
                    "style": style_name,
                    "is_heading": 'Heading' in style_name or 'Title' in style_name
                })
                
        return sections
    except Exception as e:
        logger.error(f"Failed to parse DOCX paper: {e}")
        return []


def _identify_paper_structure(sections):
    """
    识别论文结构
    
    返回结构化的章节列表：
    - title: 标题
    - abstract: 摘要
    - keywords: 关键词
    - introduction: 引言
    - body: 正文
    - conclusion: 结论
    - references: 参考文献
    """
    structure_patterns = {
        "abstract": [
            r"abstract", r"摘要", r"概要", r"summary",
            r"ABSTRACT", r"Abstract"
        ],
        "keywords": [
            r"keywords?", r"关键词", r"key words",
            r"Keywords", r"KEYWORDS"
        ],
        "introduction": [
            r"introduction", r"引言", r"前言", r"绪论",
            r"Introduction", r"INTRODUCTION"
        ],
        "conclusion": [
            r"conclusion", r"结论", r"总结", r"小结",
            r"Conclusion", r"CONCLUSIONS"
        ],
        "references": [
            r"references?", r"参考文献", r"引用文献", r"bibliography",
            r"References", r"REFERENCES"
        ]
    }
    
    result = []
    current_section = "body"
    current_content = []
    
    for section in sections:
        if isinstance(section, tuple):
            text = section[0]
        elif isinstance(section, dict):
            text = section.get("text", "")
        else:
            text = str(section)
            
        text_lower = text.lower().strip()
        
        # 检测是否为章节标题
        detected_section = None
        for section_type, patterns in structure_patterns.items():
            for pattern in patterns:
                if re.match(rf"^{pattern}", text_lower, re.IGNORECASE):
                    detected_section = section_type
                    break
            if detected_section:
                break
        
        if detected_section:
            # 保存当前章节
            if current_content and any(c.strip() for c in current_content):
                result.append((current_section, current_content))
            current_section = detected_section
            current_content = [text]
        else:
            current_content.append(text)
    
    # 保存最后一个章节
    if current_content and any(c.strip() for c in current_content):
        result.append((current_section, current_content))
    
    # 如果没有检测到任何结构，返回整体作为body
    if not result:
        all_text = []
        for section in sections:
            if isinstance(section, tuple):
                all_text.append(section[0])
            elif isinstance(section, dict):
                all_text.append(section.get("text", ""))
            else:
                all_text.append(str(section))
        result = [("body", all_text)]
    
    return result


def _identify_paper_structure_from_docx(sections):
    """从DOCX中识别论文结构"""
    structure_patterns = {
        "abstract": [r"abstract", r"摘要"],
        "keywords": [r"keywords?", r"关键词"],
        "introduction": [r"introduction", r"引言"],
        "conclusion": [r"conclusion", r"结论"],
        "references": [r"references?", r"参考文献"]
    }
    
    result = []
    current_section = "body"
    current_content = []
    
    for item in sections:
        text = item.get("text", "")
        text_lower = text.lower().strip()
        
        detected_section = None
        for section_type, patterns in structure_patterns.items():
            for pattern in patterns:
                if re.match(rf"^{pattern}", text_lower, re.IGNORECASE):
                    detected_section = section_type
                    break
            if detected_section:
                break
        
        if detected_section and item.get("is_heading"):
            if current_content:
                result.append((current_section, current_content))
            current_section = detected_section
            current_content = [text]
        else:
            current_content.append(text)
    
    if current_content:
        result.append((current_section, current_content))
    
    if not result:
        result = [("body", [item.get("text", "") for item in sections])]
    
    return result


if __name__ == "__main__":
    import sys
    
    def dummy_callback(prog=None, msg=""):
        print(f"[{prog}] {msg}" if prog else msg)
    
    if len(sys.argv) > 1:
        result = chunk(sys.argv[1], callback=dummy_callback)
        print(f"Total chunks: {len(result)}")
