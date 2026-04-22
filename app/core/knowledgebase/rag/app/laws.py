"""
法律法规文档切片方法
支持法律法规条文的结构化解析和条款提取
"""

import copy
import re
import logging
from io import BytesIO

logger = logging.getLogger(__name__)


def chunk(filename, binary=None, from_page=0, to_page=100000, lang="Chinese", callback=None, **kwargs):
    """
    法律法规切片函数
    
    支持PDF、DOCX、TXT格式的法律文件解析，按条款/章节结构进行切片
    
    Args:
        filename: 文件名或文件路径
        binary: 文件二进制数据（可选）
        from_page: 起始页码
        to_page: 结束页码
        lang: 语言 ("Chinese" 或 "English")
        callback: 进度回调函数 callback(progress, message)
        **kwargs: 其他参数
        
    Returns:
        list: 切片后的文档列表
    """
    from ..nlp import (
        naive_merge,
        tokenize_chunks,
        rag_tokenizer,
        tokenize_doc,
        is_english,
    )
    
    parser_config = kwargs.get("parser_config", {
        "chunk_token_num": 256,
        "delimiter": "\n",
        "layout_recognize": "DeepDOC",
    })
    
    eng = lang.lower() == "english"
    doc = {"docnm_kwd": filename}
    doc["title_tks"] = rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", filename))
    doc["title_sm_tks"] = rag_tokenizer.fine_grained_tokenize(doc["title_tks"])
    res = []
    
    # PDF文件处理
    if re.search(r"\.pdf$", filename, re.IGNORECASE):
        callback(0.1, "开始解析PDF法规")
        
        sections, tables, pdf_parser = _parse_laws_pdf(
            filename=filename,
            binary=binary,
            from_page=from_page,
            to_page=to_page,
            callback=callback,
            **kwargs
        )
        
        if not sections and not tables:
            return []
        
        # 提取法律条文
        law_articles = _extract_law_articles(sections)
        
        for article_num, article_title, article_content in law_articles:
            d = copy.deepcopy(doc)
            d["article_number_kwd"] = article_num
            d["article_title_kwd"] = article_title
            
            full_text = f"{article_title}\n{article_content}" if article_title else article_content
            tokenize_doc(d, full_text, eng)
            res.append(d)
        
        callback(0.8, "PDF法规解析完成")
        
    # DOCX文件处理
    elif re.search(r"\.docx$", filename, re.IGNORECASE):
        callback(0.1, "开始解析DOCX法规")
        
        sections = _parse_docx(filename, binary)
        law_articles = _extract_law_articles_from_text(sections)
        
        for article_num, article_title, article_content in law_articles:
            d = copy.deepcopy(doc)
            d["article_number_kwd"] = article_num
            d["article_title_kwd"] = article_title
            
            full_text = f"{article_title}\n{article_content}" if article_title else article_content
            tokenize_doc(d, full_text, eng)
            res.append(d)
        
        callback(0.8, "DOCX法规解析完成")
        
    # TXT文件处理
    elif re.search(r"\.txt$", filename, re.IGNORECASE):
        callback(0.1, "开始解析TXT法规")
        
        text = _read_txt_file(filename, binary)
        law_articles = _extract_law_articles_from_txt(text)
        
        for article_num, article_title, article_content in law_articles:
            d = copy.deepcopy(doc)
            d["article_number_kwd"] = article_num
            d["article_title_kwd"] = article_title
            
            full_text = f"{article_title}\n{article_content}" if article_title else article_content
            tokenize_doc(d, full_text, eng)
            res.append(d)
        
        callback(0.8, "TXT法规解析完成")
        
    else:
        raise NotImplementedError(f"不支持的文件类型: {filename} (支持: pdf, docx, txt)")
    
    return res


def _parse_laws_pdf(filename, binary=None, from_page=0, to_page=100000, callback=None, **kwargs):
    """解析PDF法规文件"""
    try:
        from app.core.knowledgebase.deepdoc.parser import PdfParser
        
        pdf_parser = PdfParser()
        sections, tables = pdf_parser(
            filename if not binary else binary, 
            from_page=from_page, 
            to_page=to_page, 
            callback=callback
        )
        return sections, tables, pdf_parser
    except Exception as e:
        logger.error(f"Failed to parse PDF laws: {e}")
        return [], [], None


def _parse_docx(filename, binary=None):
    """解析DOCX法规文件"""
    try:
        from docx import Document
        
        doc = Document(filename) if not binary else Document(BytesIO(binary))
        sections = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                sections.append(text)
                
        return sections
    except Exception as e:
        logger.error(f"Failed to parse DOCX laws: {e}")
        return []


def _read_txt_file(filename, binary=None):
    """读取TXT文件"""
    from ..nlp import find_codec
    
    if binary:
        encoding = find_codec(binary)
        return binary.decode(encoding, errors="ignore")
    else:
        with open(filename, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()


def _extract_law_articles(sections):
    """
    从解析结果中提取法律条文
    
    识别模式：
    - 第X条 / 第X章
    - Article X / Chapter X
    - 第一条 / 第一章
    """
    articles = []
    
    # 条款编号正则表达式
    article_pattern = re.compile(
        r"(第[一二三四五六七八九十百千万零〇]+[条章节])|"
        r"(第\d+[条章节])|"
        r"(Article\s*\d+)|"
        r"(Chapter\s*\d+)",
        re.IGNORECASE | re.MULTILINE
    )
    
    current_article_num = ""
    current_article_title = ""
    current_content = []
    
    for section in sections:
        if isinstance(section, tuple):
            text = section[0]
        elif isinstance(section, str):
            text = section
        else:
            continue
            
        lines = text.split('\n')
        
        for line in lines:
            line_stripped = line.strip()
            
            match = article_pattern.search(line_stripped)
            if match:
                # 保存前一个条款
                if current_content:
                    articles.append((
                        current_article_num,
                        current_article_title,
                        '\n'.join(current_content)
                    ))
                
                # 开始新条款
                current_article_num = match.group(0)
                title_part = line_stripped[match.end():].strip()
                current_article_title = title_part if title_part else ""
                current_content = [line]
            else:
                if current_article_num or articles:  # 已经开始收集条款
                    current_content.append(line)
                else:
                    # 还没遇到第一个条款，可能是标题等信息
                    pass
    
    # 保存最后一个条款
    if current_content:
        articles.append((
            current_article_num,
            current_article_title,
            '\n'.join(current_content)
        ))
    
    return articles


def _extract_law_articles_from_text(sections):
    """从文本列表中提取法律条文"""
    all_text = '\n'.join(sections)
    return _extract_law_articles_from_txt(all_text)


def _extract_law_articles_from_txt(text):
    """从纯文本中提取法律条文"""
    return _extract_law_articles([text])


if __name__ == "__main__":
    import sys
    
    def dummy_callback(prog=None, msg=""):
        print(f"[{prog}] {msg}" if prog else msg)
    
    if len(sys.argv) > 1:
        result = chunk(sys.argv[1], callback=dummy_callback)
        print(f"Total chunks: {len(result)}")
