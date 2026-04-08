"""
手册/说明书文档切片方法
支持技术手册、用户指南等文档的结构化解析
"""

import copy
import re
import logging
from io import BytesIO

logger = logging.getLogger(__name__)


def chunk(filename, binary=None, from_page=0, to_page=100000, lang="Chinese", callback=None, **kwargs):
    """
    手册切片函数
    
    支持PDF、DOCX、MD等格式的技术手册解析，按章节/步骤结构进行切片
    
    Args:
        filename: 文件名或文件路径
        binary: 文件二进制数据（可选）
        from_page: 起始页码
        to_page: 结束页码
        lang: 语言 ("Chinese" 或 "English")
        callback: 进度回调函数 callback(progress, message)
        **kwargs: 其他参数
        
    Returns:
        list: 切片后的文档列表
    """
    from ..nlp import (
        naive_merge,
        tokenize_chunks,
        rag_tokenizer,
        tokenize,
        is_english,
    )
    
    parser_config = kwargs.get("parser_config", {
        "chunk_token_num": 512,
        "delimiter": "\n",
        "layout_recognize": "DeepDOC",
    })
    
    eng = lang.lower() == "english"
    doc = {"docnm_kwd": filename}
    doc["title_tks"] = rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", filename))
    doc["title_sm_tks"] = rag_tokenizer.fine_grained_tokenize(doc["title_tks"])
    res = []
    
    # PDF文件处理
    if re.search(r"\.pdf$", filename, re.IGNORECASE):
        callback(0.1, "开始解析PDF手册")
        
        sections, tables, pdf_parser = _parse_manual_pdf(
            filename=filename,
            binary=binary,
            from_page=from_page,
            to_page=to_page,
            callback=callback,
            **kwargs
        )
        
        if not sections and not tables:
            return []
        
        # 按章节结构组织
        manual_sections = _organize_manual_sections(sections)
        
        for section_title, section_content in manual_sections:
            d = copy.deepcopy(doc)
            d["manual_section_kwd"] = section_title
            
            if isinstance(section_content, list):
                content = "\n".join(section_content)
            else:
                content = str(section_content)
                
            tokenize(d, content, eng)
            res.append(d)
        
        callback(0.8, "PDF手册解析完成")
        
    # DOCX文件处理
    elif re.search(r"\.docx$", filename, re.IGNORECASE):
        callback(0.1, "开始解析DOCX手册")
        
        sections = _parse_docx(filename, binary)
        manual_sections = _organize_manual_docx_sections(sections)
        
        for section_title, section_content in manual_sections:
            d = copy.deepcopy(doc)
            d["manual_section_kwd"] = section_title
            
            if isinstance(section_content, list):
                content = "\n".join(section_content)
            else:
                content = str(section_content)
                
            tokenize(d, content, eng)
            res.append(d)
        
        callback(0.8, "DOCX手册解析完成")
        
    # Markdown文件处理
    elif re.search(r"\.(md|markdown)$", filename, re.IGNORECASE):
        callback(0.1, "开始解析Markdown手册")
        
        text = _read_markdown(filename, binary)
        manual_sections = _organize_markdown_sections(text)
        
        for section_title, section_content in manual_sections:
            d = copy.deepcopy(doc)
            d["manual_section_kwd"] = section_title
            tokenize(d, section_content, eng)
            res.append(d)
        
        callback(0.8, "Markdown手册解析完成")
        
    else:
        raise NotImplementedError(f"不支持的文件类型: {filename} (支持: pdf, docx, md)")
    
    return res


def _parse_manual_pdf(filename, binary=None, from_page=0, to_page=100000, callback=None, **kwargs):
    """解析PDF手册"""
    try:
        from app.core.knowledge.deepdoc.parser import PdfParser
        
        pdf_parser = PdfParser()
        sections, tables = pdf_parser(
            filename if not binary else binary, 
            from_page=from_page, 
            to_page=to_page, 
            callback=callback
        )
        return sections, tables, pdf_parser
    except Exception as e:
        logger.error(f"Failed to parse PDF manual: {e}")
        return [], [], None


def _parse_docx(filename, binary=None):
    """解析DOCX手册"""
    try:
        from docx import Document
        
        doc = Document(filename) if not binary else Document(BytesIO(binary))
        sections = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                style_name = para.style.name if para.style else ""
                sections.append({
                    "text": text,
                    "style": style_name,
                    "is_heading": any(kw in style_name.lower() for kw in ['heading', 'title'])
                })
                
        return sections
    except Exception as e:
        logger.error(f"Failed to parse DOCX manual: {e}")
        return []


def _read_markdown(filename, binary=None):
    """读取Markdown文件"""
    from ..nlp import find_codec
    
    if binary:
        encoding = find_codec(binary)
        return binary.decode(encoding, errors="ignore")
    else:
        with open(filename, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()


def _organize_manual_sections(sections):
    """
    组织手册章节结构
    
    识别模式：
    - 第X章 / Chapter X
    - X. 标题 / X 标题
    - 1. / 2. / 3.
    """
    chapter_pattern = re.compile(
        r"(第[一二三四五六七八九十百千万零〇]+[章部分节])|"
        r"(第\d+[章部分节])|"
        r"(Chapter\s*\d+)|"
        r"^(\d+[\.\s]\s*.+)",
        re.MULTILINE | re.IGNORECASE
    )
    
    result = []
    current_section = ("Introduction", [])
    
    for section in sections:
        if isinstance(section, tuple):
            text = section[0]
        elif isinstance(section, str):
            text = section
        else:
            continue
            
        match = chapter_pattern.match(text.strip())
        if match:
            if current_section[1]:
                result.append(current_section)
            
            title = text.strip()
            current_section = (title, [text])
        else:
            current_section[1].append(text)
    
    if current_section[1]:
        result.append(current_section)
    
    if not result:
        all_text = []
        for s in sections:
            if isinstance(s, tuple):
                all_text.append(s[0])
            elif isinstance(s, str):
                all_text.append(s)
        result = [("Content", all_text)]
    
    return result


def _organize_manual_docx_sections(sections):
    """从DOCX中组织手册章节"""
    result = []
    current_section = ("Introduction", [])
    
    for item in sections:
        text = item.get("text", "")
        
        if item.get("is_heading") and current_section[1]:
            result.append(current_section)
            current_section = (text, [text])
        else:
            current_section[1].append(text)
    
    if current_section[1]:
        result.append(current_section)
    
    if not result:
        result = [("Content", [item.get("text", "") for item in sections])]
    
    return result


def _organize_markdown_sections(text):
    """从Markdown文本中组织章节"""
    heading_pattern = re.compile(r'^(#{1,6})\s+(.+)$', re.MULTILINE)
    
    lines = text.split('\n')
    result = []
    current_section = ("Title", "")
    current_content = []
    
    for line in lines:
        match = heading_pattern.match(line)
        if match:
            level = len(match.group(1))
            title = match.group(2).strip()
            
            # 保存当前章节（如果是H2或更高层级）
            if current_content and level >= 2:
                result.append((current_section[0], '\n'.join(current_content)))
                current_section = (title, '')
                current_content = []
            elif level == 1:
                # H1作为整体标题
                current_section = (title, '')
                current_content = []
            
            current_content.append(line)
        else:
            current_content.append(line)
    
    if current_content:
        result.append((current_section[0], '\n'.join(current_content)))
    
    if not result:
        result = [("Content", text)]
    
    return result


if __name__ == "__main__":
    import sys
    
    def dummy_callback(prog=None, msg=""):
        print(f"[{prog}] {msg}" if prog else msg)
    
    if len(sys.argv) > 1:
        result = chunk(sys.argv[1], callback=dummy_callback)
        print(f"Total chunks: {len(result)}")
